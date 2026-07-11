"""Extract text from all RareStorm PDFs/DOCX/MD into research/rarestorm-extract."""
from pathlib import Path
import sys

src = Path(r"C:\Users\MP3-Backup\Documents\GitHub\RareStorm")
out = Path(r"C:\Users\MP3-Backup\Documents\GitHub\EXO\research\rarestorm-extract")
out.mkdir(parents=True, exist_ok=True)

import fitz

for p in sorted(src.glob("*.pdf")):
    print(f"PDF: {p.name} ({p.stat().st_size} bytes)")
    doc = fitz.open(p)
    texts = []
    for i, page in enumerate(doc):
        t = page.get_text("text")
        texts.append(f"\n\n===== PAGE {i+1}/{len(doc)} =====\n{t}")
    full = "".join(texts)
    out_path = out / (p.stem + ".txt")
    out_path.write_text(full, encoding="utf-8", errors="replace")
    print(f"  pages={len(doc)} chars={len(full)} -> {out_path.name}")
    doc.close()

from docx import Document

for p in sorted(src.glob("*.docx")):
    print(f"DOCX: {p.name}")
    try:
        d = Document(str(p))
        parts = []
        for para in d.paragraphs:
            parts.append(para.text)
        for ti, table in enumerate(d.tables):
            parts.append(f"\n--- TABLE {ti+1} ---")
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
                parts.append(" || ".join(cells))
        full = "\n".join(parts)
        out_path = out / (p.stem + ".txt")
        out_path.write_text(full, encoding="utf-8", errors="replace")
        print(f"  paras={len(d.paragraphs)} tables={len(d.tables)} chars={len(full)} -> {out_path.name}")
    except Exception as e:
        print(f"  FAILED: {e}")

for p in sorted(src.glob("*.md")):
    dest = out / p.name
    dest.write_text(p.read_text(encoding="utf-8", errors="replace"), encoding="utf-8")
    print(f"MD copied: {p.name}")

print("DONE")
for f in sorted(out.glob("*.txt")):
    print(f"  {f.name}: {f.stat().st_size} bytes")
